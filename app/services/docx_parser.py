from typing import Optional
from docx import Document
from app.utils.logger import get_logger

logger = get_logger(__name__)

def extract_text_from_docx(file_path: str) -> Optional[str]:
    """
    Extract text from a DOCX file using python-docx.
    
    Args:
        file_path (str): Path to the DOCX file
        
    Returns:
        Optional[str]: Extracted text or None if extraction fails
    """
    try:
        doc = Document(file_path)
        
        # Extract text from paragraphs
        paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        
        # Extract text from tables
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    tables_text.append(" | ".join(row_text))
        
        # Combine all text
        all_text = paragraphs + tables_text
        if not all_text:
            logger.warning(f"No text content extracted from DOCX: {file_path}")
            return None
            
        return "\n\n".join(all_text)
        
    except Exception as e:
        logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
        return None 